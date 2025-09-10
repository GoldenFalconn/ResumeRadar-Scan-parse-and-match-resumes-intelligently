# extract_text.py
import pdfplumber
from docx import Document

def extract_text_from_pdf(path):
    text = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text.append(page_text)
    return "\n".join(text)

def extract_text_from_docx(path):
    doc = Document(path)
    text = [p.text for p in doc.paragraphs if p.text]
    return "\n".join(text)

# def extract_text(path):
#     if path.lower().endswith(".pdf"):
#         return extract_text_from_pdf(path)
#     elif path.lower().endswith((".docx", ".doc")):
#         return extract_text_from_docx(path)
#     else:
#         raise ValueError("Unsupported file type")

import pdfplumber
from docx import Document

def extract_text(file):
    """
    Accepts either:
    - a file path (str) 
    - a file-like object (BytesIO) from Streamlit uploader
    """
    # If it's a string (path)
    if isinstance(file, str):
        if file.lower().endswith(".pdf"):
            return extract_text_from_pdf(file)
        elif file.lower().endswith((".docx", ".doc")):
            return extract_text_from_docx(file)
        else:
            # For txt files
            with open(file, "r", encoding="utf-8") as f:
                return f.read()
    else:
        # It's a Streamlit uploaded file (BytesIO)
        name = getattr(file, "name", "")
        file.seek(0)
        if name.lower().endswith(".pdf"):
            text = []
            with pdfplumber.open(file) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text.append(page_text)
            return "\n".join(text)
        elif name.lower().endswith((".docx", ".doc")):
            doc = Document(file)
            text = [p.text for p in doc.paragraphs if p.text]
            return "\n".join(text)
        else:  # txt file
            return file.read().decode("utf-8")
